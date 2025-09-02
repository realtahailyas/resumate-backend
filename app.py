import os
from datetime import datetime
from io import BytesIO

from flask import Flask, request, jsonify
from flask_cors import CORS

# Optional: Google Generative AI (Gemini)
try:
    import google.generativeai as genai
except Exception:  # pragma: no cover
    genai = None

# File parsing deps
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:  # pragma: no cover
    pdf_extract_text = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})


# ---------------------------
# Config
# ---------------------------
DAILY_LIMIT = int(os.getenv("DAILY_LIMIT", "10"))
GENAI_API_KEY = os.getenv("GENAI_API_KEY")
GENAI_MODEL = os.getenv("GENAI_MODEL", "gemini-1.5-flash")

if genai and GENAI_API_KEY:
    try:
        genai.configure(api_key=GENAI_API_KEY)
        gemini_model = genai.GenerativeModel(GENAI_MODEL)
    except Exception:
        gemini_model = None
else:
    gemini_model = None


# ---------------------------
# In-memory store
# ---------------------------
users = {}


def _get_daily(user_id: str):
    if not user_id:
        user_id = "anon"
    if user_id not in users:
        users[user_id] = {"daily": {"date": datetime.utcnow().date().isoformat(), "count": 0}, "history": [], "cloud": []}
    rec = users[user_id]
    today = datetime.utcnow().date().isoformat()
    daily = rec.get("daily") or {}
    if daily.get("date") != today:
        rec["daily"] = {"date": today, "count": 0}
    return users[user_id]["daily"]


def can_consume(user_id: str, cost: int = 1) -> bool:
    daily = _get_daily(user_id)
    return (daily.get("count", 0) + cost) <= DAILY_LIMIT


def consume(user_id: str, cost: int = 1) -> int:
    daily = _get_daily(user_id)
    daily["count"] = daily.get("count", 0) + cost
    remaining = max(0, DAILY_LIMIT - daily["count"])
    return remaining


# ---------------------------
# Helpers
# ---------------------------
FREE_TEMPLATES = [
    {"name": "Classic", "description": "Clean, ATS-safe with clear section headers and bullet points.", "ats_friendly": True},
    {"name": "Modern", "description": "Contemporary layout using simple separators; still ATS-safe.", "ats_friendly": True},
    {"name": "Minimal", "description": "Ultra-simple text format perfect for strict ATS filters.", "ats_friendly": True},
]


def _listify_skills(skills):
    if not skills:
        return []
    if isinstance(skills, str):
        return [s.strip() for s in skills.split(",") if s.strip()]
    return [str(s).strip() for s in skills if str(s).strip()]


def _section(title):
    return f"\n{title}\n" + ("-" * len(title))


def render_resume_fallback(payload: dict, template: str = "Classic") -> str:
    full_name = payload.get("full_name") or payload.get("name") or "Your Name"
    title = payload.get("title") or "Professional Title"
    email = payload.get("email") or ""
    phone = payload.get("phone") or ""
    location = payload.get("location") or ""
    links = payload.get("links") or []
    summary = payload.get("summary") or f"Results-driven {title.lower()} with a track record of delivering impact."
    skills = _listify_skills(payload.get("skills"))
    work = payload.get("work") or payload.get("experience") or []
    education = payload.get("education") or []

    header_lines = [full_name.upper(), title]
    contact_bits = [b for b in [email, phone, location] if b]
    if links and isinstance(links, list):
        contact_bits.extend([l for l in links if l])
    if contact_bits:
        header_lines.append(" | ".join(contact_bits))

    sep = "-" * 60
    out = ["\n".join(header_lines), sep]
    out.append(_section("SUMMARY"))
    out.append(summary)

    if skills:
        out.append(_section("SKILLS"))
        out.append(" • ".join(skills))

    if work and isinstance(work, list):
        out.append(_section("EXPERIENCE"))
        for w in work:
            role = w.get("role") or w.get("title") or "Role"
            company = w.get("company") or "Company"
            period = w.get("period") or w.get("dates") or ""
            bullets = w.get("bullets") or w.get("achievements") or []
            out.append(f"{role} — {company} {('(' + period + ')') if period else ''}")
            for b in bullets:
                out.append(f"- {b}")

    if education and isinstance(education, list):
        out.append(_section("EDUCATION"))
        for e in education:
            school = e.get("school") or e.get("institution") or "School"
            degree = e.get("degree") or "Degree"
            year = e.get("year") or e.get("date") or ""
            out.append(f"{degree} — {school} {('(' + str(year) + ')') if year else ''}")

    out.append("\n")
    return "\n".join(out)


def _ai_generate(prompt: str) -> str:
    if gemini_model is None:
        return ""
    try:
        resp = gemini_model.generate_content(prompt)
        return (resp.text or "").strip()
    except Exception:
        return ""


# ---------------------------
# API Endpoints
# ---------------------------
@app.get("/api/health")
def health():
    return jsonify({"ok": True, "model": bool(gemini_model)})


@app.get("/api/credits")
def credits():
    user_id = request.args.get("user_id") or "anon"
    daily = _get_daily(user_id)
    remaining = max(0, DAILY_LIMIT - daily.get("count", 0))
    return jsonify({"credits": remaining, "credits_left": remaining, "daily_limit": DAILY_LIMIT})


@app.get("/api/templates")
def templates():
    return jsonify({"templates": FREE_TEMPLATES})


@app.get("/api/history")
def history():
    user_id = request.args.get("user_id") or "anon"
    rec = users.get(user_id, {})
    return jsonify({"history": rec.get("history", [])})


@app.post("/api/parse_jd")
def parse_jd():
    data = request.get_json(force=True, silent=True) or {}
    jd = (data.get("job_description") or "").strip()
    if not jd:
        return jsonify({"error": "job_description is required"}), 400

    prompt = (
        "Extract two sections from the job description as plain text.\n"
        "1) Key Skills: comma-separated concise keywords.\n"
        "2) Requirements: short actionable bullet points.\n\n"
        f"Job Description:\n{jd}\n\nOutput format strictly as:\n"
        "Key Skills: skill1, skill2, ...\n"
        "Requirements:\n- bullet 1\n- bullet 2"
    )
    out = _ai_generate(prompt)
    if not out:
        # Heuristic fallback
        skills = []
        reqs = []
        for line in jd.splitlines():
            lw = line.lower()
            if any(k in lw for k in ["experience", "proficient", "knowledge", "familiar", "required", "responsibilities", "requirements", "must"]):
                reqs.append(line.strip())
            # very rough skill pick
            if any(k in lw for k in ["python", "java", "javascript", "sql", "aws", "flask", "react", "ml", "ai"]):
                skills.extend([w.strip() for w in line.replace("/", ",").replace("|", ",").split(",") if w.strip()])
        skills = sorted(set([s for s in skills if len(s) <= 30]))[:20]
        reqs = ["- " + r for r in reqs[:10]]
        out = f"Key Skills: {', '.join(skills) if skills else 'See description'}\nRequirements:\n" + ("\n".join(reqs) if reqs else "- See description")
    return jsonify({"parsed": out})


@app.post("/api/ats_match")
def ats_match():
    data = request.get_json(force=True, silent=True) or {}
    user_id = data.get("user_id") or "anon"
    resume = (data.get("resume") or data.get("resume_text") or "").lower()
    jd = (data.get("job_description") or "").lower()
    if not resume or not jd:
        return jsonify({"error": "resume_text and job_description are required"}), 400

    if not can_consume(user_id):
        return jsonify({"error": "Daily limit reached. Try again tomorrow."}), 429
    remaining = consume(user_id)

    # Very simple keyword match
    tokens = set([t for t in jd.replace("/", " ").replace("|", " ").split() if t.isalpha() and len(t) > 2])
    present = [t for t in tokens if t in resume]
    missing = sorted(list(tokens - set(present)))[:20]
    match_pct = int(100 * (len(present) / max(1, len(tokens))))
    return jsonify({"match_pct": match_pct, "missing": missing, "daily_remaining": remaining})


@app.post("/api/tailor_resume")
def tailor_resume():
    data = request.get_json(force=True, silent=True) or {}
    user_id = data.get("user_id") or "anon"
    resume = (data.get("resume_text") or "").strip()
    jd = (data.get("job_description") or "").strip()
    tone = data.get("tone") or "Professional"
    if not resume or not jd:
        return jsonify({"error": "resume_text and job_description are required"}), 400

    if not can_consume(user_id):
        return jsonify({"error": "Daily limit reached. Try again tomorrow."}), 429
    remaining = consume(user_id)

    prompt = (
        f"Rewrite the following resume content in a {tone} tone to better match the job description. "
        "Keep ATS-friendly formatting, concise impact bullets, and do not hallucinate facts.\n\n"
        f"Job Description:\n{jd}\n\nResume:\n{resume}"
    )
    out = _ai_generate(prompt)
    if not out:
        out = resume + "\n\n[AI Tailored to match the job requirements]"
    return jsonify({"tailored_resume": out, "daily_remaining": remaining})


@app.post("/api/generate")
def generate():
    data = request.get_json(force=True, silent=True) or {}
    user_id = data.get("user_id") or "anon"
    jd = (data.get("job_description") or "").strip()
    resume = (data.get("resume_text") or "").strip()
    tone = data.get("tone") or "Professional"
    template = data.get("template") or "Classic"
    if not jd:
        return jsonify({"error": "job_description is required"}), 400

    if not can_consume(user_id):
        return jsonify({"error": "Daily limit reached. Try again tomorrow."}), 429
    remaining = consume(user_id)

    cover_prompt = (
        f"Write a {tone} cover letter tailored to the job description, leveraging relevant parts of the resume when provided. "
        "Keep it concise (180-250 words).\n\n"
        f"Job Description:\n{jd}\n\nResume (optional):\n{resume}"
    )
    cover = _ai_generate(cover_prompt)
    if not cover:
        cover = "Dear Hiring Manager,\n\nI am excited to apply for this role. My background aligns with your requirements across key skills and experience. " \
                "I would welcome the opportunity to contribute and help the team achieve strong outcomes.\n\nSincerely,\nYour Name"

    sugg_prompt = (
        "Provide 4-6 bullet-point suggestions to strengthen the resume for ATS based on the job description (keywords, quantifiable impact, alignment)."
        f"\n\nJob Description:\n{jd}"
    )
    suggestions = _ai_generate(sugg_prompt)
    if not suggestions:
        suggestions = "- Emphasize relevant keywords found in the job description\n- Add quantifiable results to experience bullets\n- Highlight tools and technologies explicitly\n- Keep format simple and ATS-friendly\n- Reorder bullets to show the most relevant impact first"

    users[user_id]["history"].append({
        "ts": datetime.utcnow().isoformat() + "Z",
        "action": "generate",
        "template": template,
    })

    return jsonify({"cover_letter": cover, "resume_suggestions": suggestions, "daily_remaining": remaining})


@app.post("/api/build_resume")
def build_resume():
    data = request.get_json(force=True, silent=True) or {}
    user_id = data.get("user_id") or "anon"
    template = data.get("template") or "Classic"

    # If content is sparse, optionally use AI to expand summary (costs 1)
    use_ai = bool(data.get("use_ai")) or not data.get("summary")
    if use_ai and can_consume(user_id) and gemini_model is not None:
        consume(user_id)
        base = {k: data.get(k) for k in ["title", "skills", "work", "education"]}
        prompt = (
            "Create a concise professional summary (3-4 sentences) using the candidate context."
            f"\nContext: {base}"
        )
        gen_summary = _ai_generate(prompt)
        if gen_summary:
            data["summary"] = gen_summary

    resume_text = render_resume_fallback(data, template=template)
    return jsonify({"resume_text": resume_text, "template": template})


@app.post("/api/resume_score")
def resume_score():
    data = request.get_json(force=True, silent=True) or {}
    resume = (data.get("resume_text") or "").lower()
    job_title = (data.get("job_title") or "").lower()
    jd = (data.get("job_description") or "").lower()
    if not resume:
        return jsonify({"error": "resume_text is required"}), 400
    base = resume + " " + job_title + " " + jd
    # simple heuristic
    keywords = ["python", "flask", "javascript", "sql", "api", "cloud", "aws", "react", "ml", "ai"]
    hits = sum(1 for k in keywords if k in base)
    score = min(100, 30 + hits * 7)  # rough scale
    missing = [k for k in keywords if k not in base][:10]
    return jsonify({"score": score, "missing_keywords": missing})


@app.post("/api/export")
def export_text():
    data = request.get_json(force=True, silent=True) or {}
    text = data.get("text", "")
    fmt = (data.get("format") or "txt").lower()
    instructions = {
        "pdf": "Copy into a word processor (Docs/Word) and Export > PDF.",
        "docx": "Copy into Word/Docs and Save as DOCX.",
        "txt": "Plain text ready to copy/paste.",
        "png": "Paste into a doc, format, then screenshot or print to PNG.",
    }
    return jsonify({"text": text, "format": fmt, "instructions": instructions.get(fmt, instructions["txt"])})


@app.post("/api/parse_resume_file")
def parse_resume_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    filename = (f.filename or "").lower()
    try:
        data = f.read()
        text = ""
        if filename.endswith(".pdf"):
            if pdf_extract_text is None:
                return jsonify({"error": "PDF parsing not available on server"}), 501
            text = pdf_extract_text(BytesIO(data)) or ""
        elif filename.endswith(".docx"):
            if Document is None:
                return jsonify({"error": "DOCX parsing not available on server"}), 501
            doc = Document(BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = data.decode("utf-8", errors="ignore")
        text = (text or "").strip()
        if not text:
            return jsonify({"error": "Could not extract text from file"}), 422
        return jsonify({"text": text})
    except Exception as e:
        return jsonify({"error": f"Failed to parse file: {e}"}), 500


# Cloud-style save (in-memory MVP)
@app.post("/api/cloud_save")
def cloud_save():
    data = request.get_json(force=True, silent=True) or {}
    user_id = data.get("user_id") or "anon"
    item = {k: data.get(k) for k in ["title", "template", "resume_text", "job_description"]}
    item["id"] = str(len(users[user_id]["cloud"]))
    users[user_id]["cloud"].append(item)
    return jsonify({"ok": True, "id": item["id"], "count": len(users[user_id]["cloud"])})


@app.get("/api/cloud_list")
def cloud_list():
    user_id = request.args.get("user_id") or "anon"
    items = users.get(user_id, {}).get("cloud", [])
    return jsonify({"items": [{"id": i["id"], "title": i.get("title"), "template": i.get("template")} for i in items]})


@app.get("/api/cloud_get")
def cloud_get():
    user_id = request.args.get("user_id") or "anon"
    _id = request.args.get("id")
    for i in users.get(user_id, {}).get("cloud", []):
        if i["id"] == _id:
            return jsonify({"item": i})
    return jsonify({"error": "not found"}), 404


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True)
from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import google.generativeai as genai
from datetime import datetime
from io import BytesIO
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

app = Flask(__name__)
CORS(app)

# Configure Gemini API
## Removed hardcoded GENAI_API_KEY. Only use environment variable at top of file.

# Dummy in-memory DB for MVP
users = {}

# Daily usage limit (per user) for AI-powered endpoints
DAILY_LIMIT = int(os.getenv('DAILY_LIMIT', '10'))

def _get_daily(user_id: str):
    """Ensure a daily counter exists for the user and return it."""
    if user_id not in users:
        users[user_id] = {'history': []}
    rec = users[user_id]
    today = datetime.utcnow().date().isoformat()
    daily = rec.get('daily')
    if not daily or daily.get('date') != today:
        rec['daily'] = {'date': today, 'count': 0}
    return rec['daily']

def can_consume(user_id: str, cost: int = 1) -> bool:
    daily = _get_daily(user_id)
    return (daily.get('count', 0) + cost) <= DAILY_LIMIT

def consume(user_id: str, cost: int = 1) -> int:
    daily = _get_daily(user_id)
    daily['count'] = daily.get('count', 0) + cost
    remaining = max(0, DAILY_LIMIT - daily['count'])
    return remaining

# ---------------------------
# Helpers for Resume Templates
# ---------------------------

FREE_TEMPLATES = {
    "Classic": {
        "description": "Clean, ATS-safe with clear section headers and bullet points.",
        "ats_friendly": True,
    },
    "Modern": {
        "description": "Contemporary layout using simple separators; still ATS-safe.",
        "ats_friendly": True,
    },
    "Minimal": {
        "description": "Ultra-simple text format perfect for strict ATS filters.",
        "ats_friendly": True,
    },
}


def _listify_skills(skills):
    if not skills:
        return []
    if isinstance(skills, str):
        return [s.strip() for s in skills.split(',') if s.strip()]
    return [str(s).strip() for s in skills if str(s).strip()]


def render_resume_fallback(payload: dict, template: str = "Classic") -> str:
    """Render a plain-text resume using provided fields without AI.
    Ensures ATS-friendly structure with headings and simple bullets.
    """
    full_name = payload.get('full_name') or payload.get('name') or 'Your Name'
    title = payload.get('title') or 'Professional Title'
    email = payload.get('email') or ''
    phone = payload.get('phone') or ''
    location = payload.get('location') or ''
    links = payload.get('links') or []
    summary = payload.get('summary') or f"Results-driven {title.lower()} with a track record of delivering impact."
    skills = _listify_skills(payload.get('skills'))
    work = payload.get('work') or payload.get('experience') or []
    education = payload.get('education') or []

    # Header
    header_lines = [full_name.upper(), title]
    contact_bits = [b for b in [email, phone, location] if b]
    if links and isinstance(links, list):
        contact_bits.extend([str(x) for x in links if str(x).strip()])
    if contact_bits:
        header_lines.append(' | '.join(contact_bits))

    sep = '-' * 60
    if template == 'Modern':
        section = lambda name: f"\n{name}\n" + '=' * len(name)
    elif template == 'Minimal':
        section = lambda name: f"\n{name}:"
    else:
        section = lambda name: f"\n{name}\n{sep}"

    out = []
    out.append('\n'.join(header_lines))
    out.append(sep)

    # Summary
    out.append(section('SUMMARY'))
    out.append(summary)

    # Skills
    if skills:
        out.append(section('SKILLS'))
        out.append(', '.join(skills))

    # Experience
    if work and isinstance(work, list):
        out.append(section('EXPERIENCE'))
        for job in work:
            company = job.get('company', 'Company')
            role = job.get('role') or job.get('title') or 'Role'
            start = job.get('start') or job.get('start_date') or ''
            end = job.get('end') or job.get('end_date') or 'Present'
            line = f"{role} — {company} ({start} – {end})".strip()
            out.append(line)
            ach = job.get('achievements') or job.get('bullets') or job.get('highlights') or []
            if isinstance(ach, str):
                ach = [x.strip() for x in ach.split('\n') if x.strip()]
            for a in ach[:6]:
                out.append(f"  - {a}")

    # Education
    if education and isinstance(education, list):
        out.append(section('EDUCATION'))
        for ed in education:
            school = ed.get('school') or ed.get('institution') or 'University'
            degree = ed.get('degree') or ed.get('qualification') or ''
            year = ed.get('year') or ed.get('graduation') or ''
            details = ed.get('details') or ''
            line = f"{degree} — {school} {('('+str(year)+')') if year else ''}".strip()
            out.append(line)
            if details:
                out.append(f"  - {details}")

    # Footer hint
    out.append('\n')
    return '\n'.join([x for x in out if str(x).strip()])

@app.route('/api/generate', methods=['POST'])
def generate():
    data = request.json
    user_id = data.get('user_id')
    job_description = data.get('job_description')
    resume_text = data.get('resume_text')
    tone = data.get('tone', 'Professional')
    template = data.get('template', 'Classic')

    # Enforce daily usage limit for AI calls
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    if not can_consume(user_id):
        return jsonify({'error': 'Daily limit reached. Try again tomorrow.'}), 429

    # Call Gemini API for cover letter

    try:
        cover_letter_prompt = (
            f"Write a {tone.lower()} cover letter in a {template.lower()} template for this job: {job_description} using this resume: {resume_text}"
        )
        cover_letter_resp = gemini_model.generate_content(cover_letter_prompt)
        cover_letter = cover_letter_resp.text
    except Exception as e:
        cover_letter = "[Fallback] Example cover letter: Thank you for considering my application. I am excited to apply for this role."

    try:
        resume_prompt = (
            f"Suggest resume bullet improvements for this resume: {resume_text} and job: {job_description} in a {tone.lower()} tone."
        )
        resume_resp = gemini_model.generate_content(resume_prompt)
        resume_suggestions = resume_resp.text
    except Exception as e:
        resume_suggestions = "[Fallback] Example resume bullet: Increased social media engagement by 30%."

    # Record usage and history
    remaining = consume(user_id)
    users[user_id].setdefault('history', []).append({'job': job_description, 'resume': resume_text, 'cover_letter': cover_letter})

    return jsonify({
        'cover_letter': cover_letter,
        'resume_suggestions': resume_suggestions,
        # Back-compat: expose remaining daily allowance as credits_left
        'credits_left': remaining,
        'daily_remaining': remaining,
    })

# ---------------------------
# Free AI Resume Builder
# ---------------------------
@app.route('/api/build_resume', methods=['POST'])
def build_resume():
    """Build an ATS-friendly resume from basic details. Free feature.
    Expects JSON payload with optional fields:
    {
      full_name, title, email, phone, location, links[], summary,
      skills (list or comma string),
      work: [{company, role, start, end, achievements[list or nl]}],
      education: [{school, degree, year, details}],
      template: Classic|Modern|Minimal
    }
    """
    data = request.json or {}
    template = data.get('template') or 'Classic'
    if template not in FREE_TEMPLATES:
        template = 'Classic'

    # If user provides only minimal info, we can ask AI to expand
    use_ai = False
    try:
        # Heuristic: if no achievements in work entries or summary is missing, try AI
        work = data.get('work') or []
        has_ach = any(bool((w.get('achievements') or w.get('bullets') or [])) for w in work if isinstance(w, dict))
        if not data.get('summary') or not has_ach:
            use_ai = True
    except Exception:
        use_ai = True

    # Compose AI prompt (counts against daily limit ONLY when using AI)
    if use_ai:
        user_id = (request.json or {}).get('user_id')
        if user_id and not can_consume(user_id):
            return jsonify({'error': 'Daily limit reached. Try again tomorrow.'}), 429
        try:
            # Provide a compact JSON-like representation to the model
            prompt = (
                "Create a concise, ATS-friendly resume in PLAIN TEXT (no markdown) using the following user details. "
                "Use clear section headers (SUMMARY, SKILLS, EXPERIENCE, EDUCATION). "
                "Use short, impact-driven bullet points that quantify results. "
                "Keep everything 1–2 pages max. "
                f"Style: {template}. "
                "Details:\n"
                f"Name: {data.get('full_name') or data.get('name') or 'Your Name'}\n"
                f"Title: {data.get('title') or 'Professional Title'}\n"
                f"Email: {data.get('email') or ''}\n"
                f"Phone: {data.get('phone') or ''}\n"
                f"Location: {data.get('location') or ''}\n"
                f"Links: {', '.join(data.get('links') or [])}\n"
                f"Summary (if provided): {data.get('summary') or ''}\n"
                f"Skills: {', '.join(_listify_skills(data.get('skills')))}\n"
                f"Work (company, role, dates, bullets if any): {data.get('work') or []}\n"
                f"Education: {data.get('education') or []}\n"
                "Return only the resume text, no surrounding commentary."
            )
            resp = gemini_model.generate_content(prompt)
            txt = (resp.text or '').strip()
            if txt:
                # consume only if we used AI successfully
                if user_id:
                    consume(user_id)
                return jsonify({'resume_text': txt, 'template': template})
        except Exception:
            # fall through to fallback
            pass

    # Fallback non-AI template-based rendering
    txt = render_resume_fallback(data, template)
    return jsonify({'resume_text': txt, 'template': template})

@app.route('/api/templates', methods=['GET'])
def list_templates():
    """Return available resume templates for the free library."""
    items = []
    for name, meta in FREE_TEMPLATES.items():
        items.append({
            'name': name,
            'description': meta.get('description', ''),
            'ats_friendly': bool(meta.get('ats_friendly')),
            'premium': False,
        })
    # Keep a hint that premium exists (not returned as free)
    return jsonify({'templates': items})
# ATS Match endpoint
@app.route('/api/ats_match', methods=['POST'])
def ats_match():
    data = request.json or {}
    # Accept both keys for compatibility
    resume = data.get('resume') or data.get('resume_text') or ''
    job_description = data.get('job_description', '')
    user_id = data.get('user_id')
    if user_id and not can_consume(user_id):
        return jsonify({'error': 'Daily limit reached. Try again tomorrow.'}), 429
    # Use Gemini to extract keywords from JD
    prompt = (
        "Extract a comma-separated list of the most important skills and keywords from this job description:\n"
        f"{job_description}"
    )
    try:
        resp = gemini_model.generate_content(prompt)
        jd_keywords = [kw.strip().lower() for kw in resp.text.split(',') if kw.strip()]
        resume_lower = resume.lower()
        matched = [kw for kw in jd_keywords if kw in resume_lower]
        missing = [kw for kw in jd_keywords if kw not in resume_lower]
        match_pct = int(100 * len(matched) / max(1, len(jd_keywords)))
    except Exception as e:
        match_pct = 0
        missing = []
    if user_id:
        consume(user_id)
    return jsonify({'match_pct': match_pct, 'missing': missing})

@app.route('/api/credits', methods=['GET'])
def get_credits():
    user_id = request.args.get('user_id')
    if not user_id:
        return jsonify({'credits_left': DAILY_LIMIT, 'credits': DAILY_LIMIT})
    daily = _get_daily(user_id)
    remaining = max(0, DAILY_LIMIT - daily.get('count', 0))
    # Return both keys for compatibility
    return jsonify({'credits_left': remaining, 'credits': remaining, 'daily_limit': DAILY_LIMIT})

@app.route('/api/history', methods=['GET'])
def get_history():
    user_id = request.args.get('user_id')
    if user_id not in users:
        return jsonify({'history': []})
    return jsonify({'history': users[user_id]['history']})

@app.route('/api/parse_jd', methods=['POST'])
def parse_jd():
    data = request.json
    job_description = data.get('job_description', '')
    if not job_description:
        return jsonify({'error': 'No job description provided.'}), 400
    try:
        prompt = (
            "Extract the key skills and requirements from the following job description. "
            "For skills, focus on technical skills, tools, technologies, and important soft skills. "
            "For requirements, include education, experience levels, and other qualifications. "
            "Return your answer in exactly this format:\n\n"
            "Key Skills: [comma-separated list of skills]\n"
            "Requirements: [bullet points or paragraph about requirements]\n\n"
            f"Job Description: {job_description}"
        )
        resp = gemini_model.generate_content(prompt)
        result = resp.text
        # Ensure proper formatting
        if 'Key Skills:' not in result:
            result = "Key Skills: " + result
        if 'Requirements:' not in result:
            result += "\nRequirements: See job description for details."
    except Exception as e:
        # More robust fallback: extract frequent keywords and meaningful phrases
        text = job_description.lower()
        stop = set("the a an and or for with from that this to of in on at by as you we they is are be will have has it its our your their can may must should if but not".split())
        words = [w.strip('.,:;()[]{}!?-') for w in text.split()]
        words = [w for w in words if len(w) >= 4 and w not in stop]
        freq = {}
        for w in words:
            freq[w] = freq.get(w, 0) + 1
        
        # Common technical and soft skills to look for
        common_skills = set(['python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue', 'node', 'express', 
                          'django', 'flask', 'spring', 'sql', 'nosql', 'mongodb', 'postgres', 'mysql', 'oracle',
                          'aws', 'azure', 'gcp', 'cloud', 'docker', 'kubernetes', 'devops', 'ci/cd', 'git', 
                          'agile', 'scrum', 'leadership', 'communication', 'teamwork', 'management', 'analytics',
                          'excel', 'powerpoint', 'word', 'office', 'presentation', 'sales', 'marketing',
                          'research', 'analysis', 'data', 'machine', 'learning', 'tensorflow', 'pytorch',
                          'product', 'design', 'ui/ux', 'photoshop', 'illustrator', 'figma', 'sketch'])
        
        # Identify technical skills
        technical_skills = [w for w in words if w in common_skills]
        
        # pick top 10 frequent words
        top = sorted(freq.items(), key=lambda x: (-x[1], x[0]))[:10]
        
        # Combine common skills with frequent words
        all_skills = list(set([w for w, _ in top] + technical_skills))[:15]
        skills = ', '.join(all_skills) or 'see description'
        
        # Extract education requirements
        education_phrases = ['bachelor', 'master', 'phd', 'degree', 'certification', 'diploma']
        education = []
        for phrase in education_phrases:
            if phrase in text:
                # Extract the sentence containing this phrase
                sentences = [s for s in text.split('.') if phrase in s]
                education.extend(sentences)
        
        education_text = '. '.join(education).strip() if education else 'See job description for details'
        
        result = f"Key Skills: {skills}\nRequirements: {education_text}"
    return jsonify({'parsed': result})

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'ok': True}), 200

# Tailor Resume endpoint
@app.route('/api/tailor_resume', methods=['POST'])
def tailor_resume():
    data = request.json or {}
    resume_text = data.get('resume_text', '')
    job_description = data.get('job_description', '')
    tone = data.get('tone', 'Professional')
    user_id = data.get('user_id')
    if not resume_text or not job_description:
        return jsonify({'error': 'Missing resume_text or job_description'}), 400
    if user_id and not can_consume(user_id):
        return jsonify({'error': 'Daily limit reached. Try again tomorrow.'}), 429
    try:
        prompt = (
            f"Rewrite and tailor the following resume content to better match this job description. "
            f"Keep bullet style, quantify impact where possible, and use a {tone.lower()} tone.\n\n"
            f"JOB DESCRIPTION:\n{job_description}\n\n"
            f"RESUME:\n{resume_text}\n\n"
            f"Return only the tailored resume text."
        )
        resp = gemini_model.generate_content(prompt)
        tailored = resp.text or ''
        if not tailored.strip():
            tailored = resume_text + "\n\n[AI Tailored to match the job requirements]"
        if user_id:
            consume(user_id)
        return jsonify({'tailored_resume': tailored})
    except Exception as e:
        # Fallback minimal tailoring
        fallback = resume_text + "\n\n[AI Tailored to match the job requirements]"
        return jsonify({'tailored_resume': fallback, 'warning': str(e)}), 200

# --- Cloud Save (Premium) ---
@app.route('/api/cloud_save', methods=['POST'])
def cloud_save():
    data = request.json or {}
    user_id = data.get('user_id')
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    resume_text = data.get('resume_text', '')
    job_description = data.get('job_description', '')
    template = data.get('template', 'Classic')
    title = data.get('title') or (job_description.split('\n')[0][:60] if job_description else 'Untitled')
    if user_id not in users:
        users[user_id] = {'credits': 5, 'history': []}
    user = users[user_id]
    cloud = user.get('cloud', [])
    item_id = str(len(cloud) + 1)
    item = {
        'id': item_id,
        'title': title,
        'template': template,
        'resume_text': resume_text,
        'job_description': job_description,
    }
    cloud.append(item)
    user['cloud'] = cloud
    return jsonify({'ok': True, 'id': item_id, 'count': len(cloud)})

@app.route('/api/cloud_list', methods=['GET'])
def cloud_list():
    user_id = request.args.get('user_id')
    if not user_id:
        return jsonify({'error': 'Missing user_id'}), 400
    user = users.get(user_id)
    items = (user or {}).get('cloud', [])
    # Return minimal info
    return jsonify({'items': [{'id': it['id'], 'title': it['title'], 'template': it.get('template','Classic')} for it in items]})

@app.route('/api/cloud_get', methods=['GET'])
def cloud_get():
    user_id = request.args.get('user_id')
    item_id = request.args.get('id')
    if not user_id or not item_id:
        return jsonify({'error': 'Missing parameters'}), 400
    user = users.get(user_id)
    if not user:
        return jsonify({'error': 'Not found'}), 404
    for it in user.get('cloud', []):
        if it['id'] == item_id:
            return jsonify({'item': it})
    return jsonify({'error': 'Not found'}), 404

@app.route('/api/resume_score', methods=['POST'])
def resume_score():
    """Score a resume against a job title or job description.
    Returns a percentage match and areas for improvement.
    """
    data = request.json or {}
    resume_text = data.get('resume_text', '')
    job_title = data.get('job_title', '')
    job_description = data.get('job_description', '')
    
    if not resume_text:
        return jsonify({'error': 'Missing resume_text'}), 400
    if not (job_title or job_description):
        return jsonify({'error': 'Missing job_title or job_description'}), 400
    
    # Try AI-based score if both resume and job details exist
    if resume_text and (job_title or job_description):
        try:
            target = job_description if job_description else f"a {job_title} position"
            prompt = (
                "You are an expert ATS (Applicant Tracking System) resume analyst. "
                "Score this resume for how well it matches the target job. "
                "Provide a percentage score (0-100) and identify 3-5 specific keywords or skills missing "
                "that would improve the score. Also provide 2-3 brief, specific suggestions to improve the resume.\n\n"
                f"TARGET JOB: {target}\n\n"
                f"RESUME:\n{resume_text}\n\n"
                "Return ONLY a JSON object with these fields: score (number), missing_keywords (list of strings), "
                "suggestions (list of strings). No other text."
            )
            resp = gemini_model.generate_content(prompt)
            
            try:
                # Try to parse as JSON
                import json
                import re
                
                # Extract JSON from response if needed
                text = resp.text.strip()
                match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
                if match:
                    text = match.group(1)
                elif text.startswith('{') and text.endswith('}'):
                    # Already JSON-like
                    pass
                else:
                    # Not proper JSON, try simple parsing
                    result = {}
                    if 'score' in text.lower():
                        score_match = re.search(r'score["\']?\s*:\s*(\d+)', text, re.IGNORECASE)
                        if score_match:
                            result['score'] = int(score_match.group(1))
                    return jsonify(result)
                
                result = json.loads(text)
                # Make sure score is an integer
                if 'score' in result and not isinstance(result['score'], int):
                    try:
                        result['score'] = int(float(str(result['score']).replace('%', '')))
                    except:
                        result['score'] = 70  # fallback
                
                return jsonify(result)
            except Exception as e:
                # If JSON parsing fails, extract score using regex
                score_match = re.search(r'(\d+)%', resp.text)
                if score_match:
                    score = int(score_match.group(1))
                    return jsonify({'score': score})
                else:
                    # Fallback to simplified heuristic score
                    pass
        except Exception as e:
            print(f"AI scoring error: {e}")
            # Fall through to heuristic
    
    # Fallback to simplified heuristic scoring
    keywords = []
    
    # Extract keywords from job title using common skills
    common_skills = {
        "engineer": ["development", "programming", "problem-solving", "technical", "design", "algorithms"],
        "developer": ["coding", "software", "applications", "frameworks", "github", "testing"],
        "manager": ["leadership", "team", "strategy", "communication", "project", "stakeholders"],
        "analyst": ["data", "analysis", "reporting", "insights", "excel", "visualization"],
        "designer": ["design", "user experience", "ui", "creative", "visual", "portfolio"],
        "marketing": ["campaigns", "social media", "content", "analytics", "strategy", "audience"],
        "sales": ["revenue", "clients", "relationship", "targets", "negotiation", "customer"],
    }
    
    for role, skills in common_skills.items():
        if role.lower() in job_title.lower():
            keywords.extend(skills)
    
    # Add generic professional keywords
    keywords.extend(["experience", "skills", "professional", "results", "achievements", "education"])
    
    # Calculate match
    resume_lower = resume_text.lower()
    matches = sum(1 for kw in keywords if kw.lower() in resume_lower)
    total = len(keywords)
    match_pct = min(95, int(100 * matches / max(1, total)))
    
    # Basic missing keywords
    missing = [kw for kw in keywords if kw.lower() not in resume_lower]
    
    return jsonify({
        'score': match_pct,
        'missing_keywords': missing[:5]
    })

@app.route('/api/export', methods=['POST'])
def export_document():
    """Convert resume text to various formats (PDF, DOCX, etc).
    For the MVP, we'll return plain text with formatting suggestions.
    In a real implementation, this would use a library like python-docx or reportlab.
    """
    data = request.json or {}
    text = data.get('text', '')
    format = data.get('format', 'txt').lower()
    title = data.get('title', 'Resume')
    
    if not text:
        return jsonify({'error': 'Missing text content'}), 400
    
    # For MVP, just return the text with suggestions based on format
    instructions = {
        'pdf': "To convert to PDF: Copy this text, paste into a word processor like Google Docs or Microsoft Word, format as desired, then export/save as PDF.",
        'docx': "To convert to DOCX: Copy this text, paste into Microsoft Word or Google Docs, format as desired, then save as DOCX.",
        'txt': "Plain text format ready for copy/paste.",
        'png': "To create a PNG: Copy this text into a word processor, format as desired, then use a screenshot tool or Save as PNG option."
    }
    
    return jsonify({
        'text': text,
        'format': format,
        'instructions': instructions.get(format, instructions['txt'])
    })

@app.route('/api/parse_resume_file', methods=['POST'])
def parse_resume_file():
    """Parse uploaded resume files (PDF, DOCX, or plain text).
    Accepts multipart/form-data with 'file'. Returns extracted text.
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['file']
    filename = f.filename.lower()
    try:
        data = f.read()
        text = ''
        if filename.endswith('.pdf'):
            text = pdf_extract_text(BytesIO(data)) or ''
        elif filename.endswith('.docx'):
            doc = Document(BytesIO(data))
            text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            # treat as text
            text = data.decode('utf-8', errors='ignore')
        text = (text or '').strip()
        if not text:
            return jsonify({'error': 'Could not extract text from file'}), 422
        return jsonify({'text': text})
    except Exception as e:
        return jsonify({'error': f'Failed to parse file: {e}'}), 500

if __name__ == '__main__':
    app.run(debug=True)
